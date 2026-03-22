import os
import io
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from groq import Groq
from docx import Document
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app) 

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# ai_synonyms route
@app.route('/api/ai/synonyms', methods=['POST'])
def ai_synonyms():
    data = request.json
    word = data.get('word')
    context = data.get('context')
    
    prompt = f"Give me 3 synonyms for the word '{word}' suitable for this context: '{context}'. Return only the words separated by commas."
    try:
        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[{"role": "user", "content": prompt}]
        )
        synonyms_text = response.choices[0].message.content
        synonyms = [s.strip() for s in synonyms_text.split(',')]
        return jsonify(synonyms)
    except Exception as e:
        return jsonify(["Error fetching synonyms"])
    
# ai_critique route
@app.route('/api/ai/critique', methods=['POST'])
def ai_critique():
    data = request.json
    text = data.get('text')
    history = data.get('history', []) # pull context history
    
    system_prompt = """You are an expert editor. Analyze the chapter text for flaws, inconsistencies, and pacing. 
    Format your response in Markdown. At the end, you MUST provide actionable snippets in this exact format:
    **Suggested Snippets:**
    * Change "[old text]" to "[new text]" """
    
    messages = [{"role": "system", "content": system_prompt}]
    
    # append previous chat history for context
    for msg in history:
        messages.append(msg)
        
    messages.append({"role": "user", "content": f"Text to analyze: {text[:4000]}..."})
    
    try:
        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=messages
        )
        return jsonify({"critique": response.choices[0].message.content})
    except Exception as e:
        return jsonify({"critique": str(e)})
        
@app.route('/export', methods=['POST'])
def export_docx():
    data = request.json
    doc = Document()
    doc.add_heading('Manuscript Export', 0)

    for chap in data.get('chapters', []):
        doc.add_heading(chap.get('title', 'Untitled'), level=1)
        doc.add_paragraph(chap.get('content', ''))
        doc.add_page_break()

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(file_stream, as_attachment=True, download_name='Manuscript.docx')

if __name__ == '__main__':
    app.run(debug=True, port=5000)
